import os
import re
from collections import defaultdict
from zipfile import BadZipFile

import docx
import docx2txt
import pandas as pd

from .settings import (
    MOST_FREQUENT_VALUES,
    SUBMISSION_COLUMNS,
    TABLE_PATTERNS,
    TEXT_PATTERNS
)


def get_text_from_docx(file):
    text = None
    try:
        text = docx2txt.process(file)
    except (FileNotFoundError, BadZipFile, KeyError, RuntimeError) as e:
        print(e)
        print(file)
    return text

def convert_table_to_df(table):
    try:
        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)

        df = pd.DataFrame(data)
        return df
    except IndexError as e:
        print(e)
        return None

def extract_target_from_text(text, phrase_pattern, target_pattern):
    match = re.search(phrase_pattern, text)
    if match:
        target = re.search(target_pattern, match.group())
        if target:
            return target.group()
        else:
            return None
    else:
        return None

def get_data_from_docx(file, pattern):
    try:
        f = open(file, 'rb')
        document = docx.Document(f)
        for table in document.tables:
            df = convert_table_to_df(table)
            if df is not None:
                if pattern in df.columns:
                    f.close()
                    return df[pattern].loc[0]
        f.close()
        return None
    except KeyError as e:
        print(e)
        print(file.name)
        return None

def predict_partially(project):
    df_dict = {
        column: [] for column in SUBMISSION_COLUMNS
    }
    well_clusters = []
    values = defaultdict(list)

    for document in os.scandir(project.path):

        # Извлечение данных из текста
        text = get_text_from_docx(document.path)
        if text is not None:
            text = text.lower()
            samples = re.findall(
                pattern=TEXT_PATTERNS['Куст']['phrase_pattern'],
                string=text
            )
            for sample in set(samples):
                well_clusters.extend(
                    re.findall(
                        TEXT_PATTERNS['Куст']['target_pattern'],
                        sample
                    )
                )
            for target in TEXT_PATTERNS.keys():
                if target != 'Куст':
                    value = extract_target_from_text(
                        text,
                        TEXT_PATTERNS[target]['phrase_pattern'],
                        TEXT_PATTERNS[target]['target_pattern']
                    )
                    values[target].append(value)

        # Извлечение данных из таблиц
        try:
            f = open(document, 'rb')
            doc = docx.Document(f)
            for table in doc.tables:
                df = convert_table_to_df(table)
                if df is not None:
                    for target in TABLE_PATTERNS.keys():
                        if target == 'Куст':
                            pattern = TABLE_PATTERNS['Куст']
                            if pattern in df.columns:
                                well_clusters.append(df[pattern].loc[0])
                        else:
                            pass
            f.close()
        except KeyError as e:
            print(e)
            print(document.name)


    for target in TEXT_PATTERNS.keys():
        values[target] = [x for x in values[target] if x is not None]
        if len(values[target]) > 0:
            values[target] = values[target][0]
            if target == 'Абсолютный минимум температуры':
                values[target] = - int(values[target])
        else:
            values[target] = None

    well_clusters = set(well_clusters)
    if len(well_clusters) > 0:
        for well_cluster in set(well_clusters):
            df_dict['Проект'].append(project.name)
            df_dict['Куст'].append(well_cluster)
            for column in SUBMISSION_COLUMNS[2:]:
                if column in values.keys():
                    df_dict[column].append(values[column])
                else:
                    df_dict[column].append(MOST_FREQUENT_VALUES[column])
        return pd.DataFrame(df_dict)
    else:
        return None
